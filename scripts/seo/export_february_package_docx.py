#!/usr/bin/env python3
"""
Export the February package markdown drafts to Google-Docs-friendly .docx files.

Why this exists:
- Pasting Markdown/HTML into Google Docs is inconsistent.
- .docx import preserves real Heading styles and lists.

Outputs:
  work/seo/hogeye/february_package/google_docs_ready_docx_v2/*.docx
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
FEB_DIR = REPO_ROOT / "work/seo/hogeye/february_package"
OUT_DIR = FEB_DIR / "google_docs_ready_docx_v2"

FILES = [
    "blog_01_wild_hog_trap_release_camera_system.md",
    "blog_02_sounder_capture_timing.md",
    "blog_03_trap_gate_trigger_reliability.md",
    "blog_04_multi_user_trap_monitoring.md",
    "blog_05_trap_compatibility_and_warranty.md",
]

FM_KEY_RE = re.compile(r"^([A-Za-z0-9_\-]+):\s*(.*)$")


def parse_frontmatter(text: str) -> Tuple[Dict[str, object], str]:
    if not text.startswith("---"):
        return {}, text
    parts = text.split("---", 2)
    if len(parts) < 3:
        return {}, text

    fm_lines = parts[1].strip().splitlines()
    body = parts[2].lstrip("\n")

    meta: Dict[str, object] = {}
    cur_list_key: str | None = None

    for line in fm_lines:
        if not line.strip():
            continue
        m = FM_KEY_RE.match(line)
        if m:
            k = m.group(1).strip()
            v = m.group(2).strip()
            if v == "" or v.endswith(":"):
                meta[k] = []
                cur_list_key = k
            else:
                meta[k] = v.strip('"')
                cur_list_key = None
            continue
        if line.strip().startswith("-") and cur_list_key:
            assert isinstance(meta[cur_list_key], list)
            meta[cur_list_key].append(line.strip()[1:].strip().strip('"'))

    return meta, body


def _set_style(doc: Document, style_name: str, *, size_pt: int, bold: bool) -> None:
    st = doc.styles[style_name]
    st.font.name = "Arial"
    st.font.size = Pt(size_pt)
    st.font.bold = bold
    st.font.color.rgb = RGBColor(0, 0, 0)
    # Readable, Docs-friendly spacing (avoid both "airy" and "jammed").
    st.paragraph_format.space_before = Pt(12) if style_name in {"Heading 1", "Heading 2"} else Pt(6)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    st.paragraph_format.line_spacing = 1.15


def apply_clean_styles(doc: Document) -> None:
    # Narrow page margins to reduce horizontal whitespace.
    # Google Docs imports .docx margins faithfully.
    for s in doc.sections:
        s.left_margin = Inches(0.65)
        s.right_margin = Inches(0.65)
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)

    # Normal text
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    # ChatGPT-like readable default: 1.15 with small paragraph spacing.
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # Headings: force black (Google Docs often imports docx headings as blue otherwise)
    _set_style(doc, "Heading 1", size_pt=20, bold=True)
    _set_style(doc, "Heading 2", size_pt=14, bold=True)
    _set_style(doc, "Heading 3", size_pt=12, bold=True)

    # Lists: tighten spacing as well
    for name in ("List Bullet", "List Number"):
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = "Arial"
            st.font.size = Pt(11)
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.paragraph_format.space_before = Pt(0)
            # Slightly tighter for bullets so lists don't look double-spaced.
            st.paragraph_format.space_after = Pt(2)
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            st.paragraph_format.line_spacing = 1.15


def human_label(k: str) -> str:
    k = (k or "").strip().replace("_", " ")
    return k[:1].upper() + k[1:]


def add_metadata_block(doc: Document, meta: Dict[str, object]) -> None:
    doc.add_heading("Metadata", level=2)

    def add_field(key: str, value: object) -> None:
        if value is None:
            return
        if isinstance(value, list):
            if not value:
                return
            p = doc.add_paragraph()
            r = p.add_run(f"{human_label(key)}:")
            r.bold = True
            for item in value:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            s = str(value).strip()
            if not s:
                return
            p = doc.add_paragraph()
            r = p.add_run(f"{human_label(key)}: ")
            r.bold = True
            p.add_run(s)

    # Clean, review-relevant subset (no tables)
    add_field("slug", meta.get("slug"))
    add_field("focus_keyword", meta.get("focus_keyword"))
    add_field("additional_keywords", meta.get("additional_keywords"))
    add_field("meta_title", meta.get("meta_title"))
    add_field("meta_description", meta.get("meta_description"))
    add_field("category", meta.get("category"))
    add_field("internal_links", meta.get("internal_links"))
    add_field("schema", meta.get("schema"))


def add_runs_with_bold(paragraph, text: str) -> None:
    """
    Very small markdown bold handler for **bold**.
    """
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def md_body_to_doc(doc: Document, body: str, *, title_to_drop: str | None) -> None:
    lines = body.splitlines()
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1

    # Drop a leading title heading if it duplicates the document title.
    if i < len(lines):
        m = re.match(r"^#{1,6}\s+(.*)$", lines[i].strip())
        if m and title_to_drop and m.group(1).strip() == title_to_drop.strip():
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1

    for raw in lines[i:]:
        line = raw.rstrip().replace("`", "")
        # Do not insert empty paragraphs for blank lines.
        # We keep spacing tight and let Docs/Word handle paragraph breaks naturally.
        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        if re.match(r"^\s*-\s+", line):
            txt = re.sub(r"^\s*-\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, txt.replace("**", "**"))
            continue

        if re.match(r"^\s*\d+\)\s+", line):
            txt = re.sub(r"^\s*\d+\)\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, txt)
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, line)


def export_one(md_path: Path) -> Path:
    text = md_path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)
    title = str(meta.get("title") or md_path.stem)

    doc = Document()
    apply_clean_styles(doc)

    doc.add_heading(title, level=1)
    add_metadata_block(doc, meta)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Draft", level=2)
    md_body_to_doc(doc, body, title_to_drop=title)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / (md_path.stem + ".docx")
    doc.save(out_path)
    return out_path


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for fn in FILES:
        p = FEB_DIR / fn
        if not p.is_file():
            raise SystemExit(f"Missing file: {p}")
        out = export_one(p)
        print(f"Wrote {out}")
    print(f"Done. Output dir: {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

